#!/usr/bin/env python3
"""House style for Opportunity Briefs (.docx) — the ONLY road into Google Docs.

Rules (learned the hard way):
  * Word (.docx) imports into Google Docs with styling intact. Raw HTML does NOT —
    never push HTML at Drive. Markdown paste mangles tables (literal **). docx only.
  * Default python-docx output is not acceptable: the brief is marketing as much
    as substance. Use this module so every brief shares one design system.
  * Adapt PALETTE to the target company's brand (title/accent colours) per run.

Usage (from a synthesise session):
    import importlib.util, sys
    spec = importlib.util.spec_from_file_location("brief_style", "<path>/brief-style.py")
    bs = importlib.util.module_from_spec(spec); spec.loader.exec_module(bs)

    d = bs.BriefDoc(chip="DRAFT v1 — FOR REVIEW",
                    title="Company: The opportunity",
                    meta=[("Author", "Name"), ("Contributor", "Claude"), ("Date", "..."), ("Time-box", "90 minutes")])
    d.h1("Summary")
    d.callout([("Bold thesis lead. ", True), ("Rest of the BLUF...", False)], kind="primary")
    d.h2("Barriers")
    d.para([("Barrier 1 — name. ", True), ("Body text...", False)])
    d.quote("“Customer voice quote.”")
    d.table(header=("Col A", "Col B"), rows=[("a", "b")], widths_in=(2.0, 4.0))
    d.bullets([[("Now — thing. ", True), ("Detail. ", False)]])
    d.sources("Sources: ...")
    d.save("/path/out.docx")

Run `python3 brief-style.py` to generate a styled demo at ./brief-style-demo.docx.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Adapt per company. Defaults: Checkatrade-flavoured navy/red.
PALETTE = {
    "primary": "15224F",   # headings, callout bars, table header
    "accent": "E03C31",    # chip, secondary callout bars
    "grey": "5F5F5F",
    "tint_primary": "EEF2F8",
    "tint_grey": "F5F4F1",
    "hairline": "D9D6CE",
}

def _rgb(hexstr):
    return RGBColor(int(hexstr[0:2], 16), int(hexstr[2:4], 16), int(hexstr[4:6], 16))


class BriefDoc:
    def __init__(self, chip, title, meta, palette=None):
        self.pal = dict(PALETTE, **(palette or {}))
        self.doc = Document()
        self._base_styles()
        self._title_block(chip, title, meta)

    # ---------- internals ----------
    def _base_styles(self):
        n = self.doc.styles["Normal"]
        n.font.name = "Arial"; n.font.size = Pt(10.5)
        n.paragraph_format.space_after = Pt(6)
        n.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        n.paragraph_format.line_spacing = 1.15
        for name, size, before in (("Title", 24, 0), ("Heading 1", 15, 18), ("Heading 2", 12, 12)):
            s = self.doc.styles[name]
            s.font.name = "Arial"; s.font.size = Pt(size); s.font.bold = True
            s.font.color.rgb = _rgb(self.pal["primary"])
            s.paragraph_format.space_before = Pt(before)
            s.paragraph_format.space_after = Pt(6 if name != "Title" else 2)

    def _border(self, para, edge, color, size="12", space="4"):
        pr = para._p.get_or_add_pPr()
        pbdr = pr.find(qn("w:pBdr"))
        if pbdr is None:
            pbdr = OxmlElement("w:pBdr"); pr.append(pbdr)
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), size)
        el.set(qn("w:space"), space); el.set(qn("w:color"), color)
        pbdr.append(el)

    def _shade(self, para, hexcolor):
        pr = para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hexcolor)
        pr.append(shd)

    def _runs(self, para, pieces):
        for piece in pieces:
            text, bold = piece[0], piece[1]
            italic = piece[2] if len(piece) > 2 else False
            r = para.add_run(text); r.bold = bold; r.italic = italic
        return para

    def _title_block(self, chip, title, meta):
        c = self.doc.add_paragraph()
        r = c.add_run(f"  {chip}  ")
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        self._shade(c, self.pal["accent"]); c.paragraph_format.space_after = Pt(2)
        self.doc.add_paragraph(title, style="Title")
        m = self.doc.add_paragraph()
        for label, val in meta:
            r = m.add_run(f"{label}  "); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = _rgb(self.pal["grey"])
            r = m.add_run(f"{val}     "); r.font.size = Pt(9)
        self._border(m, "bottom", self.pal["primary"], size="16", space="6")
        m.paragraph_format.space_after = Pt(14)

    # ---------- public API ----------
    def h1(self, text): self.doc.add_heading(text, 1)
    def h2(self, text): self.doc.add_heading(text, 2)

    def para(self, pieces, style=None):
        return self._runs(self.doc.add_paragraph(style=style), pieces)

    def bullets(self, items):
        for pieces in items:
            self._runs(self.doc.add_paragraph(style="List Bullet"), pieces)

    def callout(self, pieces, kind="primary"):
        """Shaded box with a coloured left bar. kind: primary | warn (accent bar, grey tint)."""
        para = self._runs(self.doc.add_paragraph(), pieces)
        if kind == "primary":
            self._shade(para, self.pal["tint_primary"]); bar = self.pal["primary"]
        else:
            self._shade(para, self.pal["tint_grey"]); bar = self.pal["accent"]
        self._border(para, "left", bar, size="24", space="8")
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(10)
        return para

    def quote(self, text, lead="Potential customer: "):
        para = self.doc.add_paragraph()
        r = para.add_run(lead); r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = _rgb(self.pal["grey"])
        r = para.add_run(text); r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = _rgb(self.pal["grey"])
        self._border(para, "left", "C9C6BE", size="18", space="8")
        para.paragraph_format.space_after = Pt(10)
        return para

    def table(self, header, rows, widths_in):
        data = [header] + list(rows)
        t = self.doc.add_table(rows=len(data), cols=len(header))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, row in enumerate(data):
            for j, text in enumerate(row):
                cell = t.cell(i, j); cell.width = Inches(widths_in[j])
                cell.text = ""
                run = cell.paragraphs[0].add_run(str(text)); run.font.size = Pt(9)
                tcPr = cell._tc.get_or_add_tcPr()
                if i == 0:
                    run.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:fill"), self.pal["primary"]); tcPr.append(shd)
                else:
                    if i % 2 == 0:
                        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
                        shd.set(qn("w:fill"), self.pal["tint_grey"]); tcPr.append(shd)
                    borders = OxmlElement("w:tcBorders")
                    for edge in ("top", "bottom"):
                        el = OxmlElement(f"w:{edge}")
                        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "4")
                        el.set(qn("w:color"), self.pal["hairline"])
                        borders.append(el)
                    tcPr.append(borders)
                mar = OxmlElement("w:tcMar")
                for side, w in (("top", "60"), ("bottom", "60"), ("left", "100"), ("right", "100")):
                    el = OxmlElement(f"w:{side}"); el.set(qn("w:w"), w); el.set(qn("w:type"), "dxa")
                    mar.append(el)
                tcPr.append(mar)
        self.doc.add_paragraph()
        return t

    def sources(self, text):
        para = self.doc.add_paragraph()
        r = para.add_run(text); r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = _rgb(self.pal["grey"])
        self._border(para, "top", self.pal["hairline"], size="8", space="6")
        return para

    def save(self, path):
        self.doc.save(path)
        return path


if __name__ == "__main__":
    d = BriefDoc(chip="DEMO — HOUSE STYLE", title="Example: The opportunity brief",
                 meta=[("Author", "Your Name"), ("Contributor", "Claude"), ("Date", "Today"), ("Time-box", "90 minutes")])
    d.h1("Summary")
    d.callout([("The thesis in bold. ", True), ("Then the supporting argument in regular weight, all in one skimmable box.", False)])
    d.h2("A section")
    d.para([("Lead-in. ", True), ("Body text with a ", False), ("bolded", True), (" middle.", False)])
    d.quote("“A customer voice quote sits in a grey left-barred block.”")
    d.table(header=("Item", "Detail", "Score"), rows=[("One", "Banded rows, hairline borders", "12.0"), ("Two", "Header row in primary colour", "3.5")], widths_in=(1.5, 3.5, 1.0))
    d.bullets([[("Now — ", True), ("do the cheap thing.", False)], [("Next — ", True), ("do the strategic thing.", False)]])
    d.sources("Sources: demo.")
    print("Saved:", d.save("brief-style-demo.docx"))
